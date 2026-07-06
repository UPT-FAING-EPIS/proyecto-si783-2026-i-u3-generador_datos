"""
Script para generar el Diccionario de Datos del Generador de Datos Sintéticos.
Reemplaza el documento de plantilla con el contenido real del proyecto.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"c:\Users\Mariela\Downloads\DATA-GENERATOR\docs\Diccionario_de_Datos_GeneradorDatos.docx"

# ──────────────────────────────────────────────────────────
# DATOS DEL PROYECTO
# ──────────────────────────────────────────────────────────
INTEGRANTES = [
    "Ramos Loza, Mariela Estefany (2023077478)",
    "Llanos Niño, Vincenzo Rafael (202307679)",
]
DOCENTE = "Ing. Patrick Cuadros"
CURSO = "Base de Datos II"
PROYECTO = "Sistema Generador de Datos Sintéticos con IA"

# Definición de las tablas reales del sistema
TABLAS = [
    {
        "nombre": "usuarios",
        "descripcion": "Tabla principal que almacena la información de registro y autenticación de cada usuario del sistema.",
        "objetivo": "Gestionar la identidad, credenciales y estado de los usuarios que utilizan la plataforma generadora de datos.",
        "relaciones": "Es referenciada por las tablas sesiones, logs, conexiones y comentarios a través de la clave foránea usuario_id.",
        "campos": [
            ("1", "id", "INTEGER", "No", "Sí", "No", "Identificador único autoincremental del usuario."),
            ("2", "nombre", "VARCHAR(100)", "No", "No", "No", "Nombre de pila del usuario."),
            ("3", "apellido", "VARCHAR(100)", "No", "No", "No", "Apellido(s) del usuario."),
            ("4", "email", "VARCHAR(255)", "No", "No", "No", "Correo electrónico único e indexado para login."),
            ("5", "password_hash", "VARCHAR(255)", "Sí", "No", "No", "Hash de la contraseña. Nulo si el acceso es exclusivamente por OAuth."),
            ("6", "rol", "ENUM", "No", "No", "No", "Rol en el sistema: 'superadmin' o 'usuario'."),
            ("7", "activo", "BOOLEAN", "No", "No", "No", "Indica si la cuenta está habilitada (True) o suspendida (False)."),
            ("8", "avatar_url", "VARCHAR(500)", "Sí", "No", "No", "URL a la imagen de perfil del usuario."),
            ("9", "created_at", "DATETIME", "No", "No", "No", "Fecha y hora de creación del registro."),
            ("10", "updated_at", "DATETIME", "No", "No", "No", "Fecha y hora de la última modificación del registro."),
            ("11", "ultimo_acceso", "DATETIME", "Sí", "No", "No", "Fecha y hora del último inicio de sesión exitoso."),
            ("12", "ultima_ip", "VARCHAR(45)", "Sí", "No", "No", "Dirección IP del último acceso (soporta IPv4 e IPv6)."),
        ]
    },
    {
        "nombre": "sesiones",
        "descripcion": "Registro de las sesiones activas e históricas de los usuarios en el sistema.",
        "objetivo": "Controlar y auditar los tokens JWT emitidos, garantizando la seguridad y el control de acceso activo.",
        "relaciones": "Referencia a la tabla usuarios mediante la clave foránea usuario_id con eliminación en cascada.",
        "campos": [
            ("1", "id", "INTEGER", "No", "Sí", "No", "Identificador único autoincremental de la sesión."),
            ("2", "usuario_id", "INTEGER", "No", "No", "Sí → usuarios.id", "ID del usuario propietario de la sesión."),
            ("3", "token_jwt", "TEXT", "No", "No", "No", "Token JWT emitido para autenticar las peticiones del usuario."),
            ("4", "ip_address", "VARCHAR(45)", "Sí", "No", "No", "Dirección IP desde donde se inició la sesión."),
            ("5", "metodo_login", "ENUM", "No", "No", "No", "Método de autenticación usado: 'email', 'google', 'github' o 'microsoft'."),
            ("6", "activa", "BOOLEAN", "No", "No", "No", "Indica si el token de sesión sigue siendo válido."),
            ("7", "created_at", "DATETIME", "No", "No", "No", "Fecha y hora de creación de la sesión."),
            ("8", "expires_at", "DATETIME", "No", "No", "No", "Fecha y hora de expiración del token JWT."),
        ]
    },
    {
        "nombre": "logs",
        "descripcion": "Tabla de auditoría que registra todas las acciones significativas realizadas dentro del sistema.",
        "objetivo": "Proveer un historial trazable de operaciones para fines de seguridad, depuración y monitoreo del sistema.",
        "relaciones": "Referencia a la tabla usuarios mediante usuario_id. En caso de eliminación del usuario, el campo se pone en NULL (SET NULL).",
        "campos": [
            ("1", "id", "INTEGER", "No", "Sí", "No", "Identificador único autoincremental del registro de log."),
            ("2", "usuario_id", "INTEGER", "Sí", "No", "Sí → usuarios.id", "ID del usuario que realizó la acción. Puede ser nulo para acciones del sistema."),
            ("3", "accion", "VARCHAR(100)", "No", "No", "No", "Nombre corto y descriptivo de la acción ejecutada."),
            ("4", "detalle", "TEXT", "Sí", "No", "No", "Información técnica adicional o descripción extendida del evento."),
            ("5", "ip_address", "VARCHAR(45)", "Sí", "No", "No", "Dirección IP desde la cual se realizó la acción registrada."),
            ("6", "nivel", "VARCHAR(20)", "No", "No", "No", "Nivel de severidad del evento: 'INFO', 'WARNING' o 'ERROR'."),
            ("7", "created_at", "DATETIME", "No", "No", "No", "Fecha y hora exacta en que se registró el evento."),
        ]
    },
    {
        "nombre": "conexiones",
        "descripcion": "Historial y configuración de las bases de datos externas a las que cada usuario se ha conectado para inyectar datos sintéticos.",
        "objetivo": "Almacenar de forma segura las credenciales de acceso a bases de datos externas y llevar estadísticas de uso de la herramienta de inyección.",
        "relaciones": "Referencia a la tabla usuarios mediante usuario_id con eliminación en cascada.",
        "campos": [
            ("1", "id", "INTEGER", "No", "Sí", "No", "Identificador único autoincremental de la conexión."),
            ("2", "usuario_id", "INTEGER", "No", "No", "Sí → usuarios.id", "ID del usuario propietario de esta configuración de conexión."),
            ("3", "nombre_alias", "VARCHAR(100)", "Sí", "No", "No", "Nombre amigable o alias asignado por el usuario a esta conexión."),
            ("4", "motor_bd", "VARCHAR(50)", "No", "No", "No", "Motor de base de datos destino: 'mysql', 'postgresql', 'mongodb', etc."),
            ("5", "host", "VARCHAR(255)", "No", "No", "No", "Dirección del servidor de base de datos (IP o dominio)."),
            ("6", "puerto", "INTEGER", "No", "No", "No", "Puerto de conexión al servidor de base de datos."),
            ("7", "nombre_bd", "VARCHAR(255)", "No", "No", "No", "Nombre de la base de datos a la que se realiza la conexión."),
            ("8", "usuario_db", "VARCHAR(255)", "Sí", "No", "No", "Nombre de usuario con acceso a la base de datos externa."),
            ("9", "password_db", "TEXT", "Sí", "No", "No", "Contraseña cifrada con Fernet para la base de datos externa."),
            ("10", "registros_generados", "INTEGER", "No", "No", "No", "Contador total de registros sintéticos generados para esta conexión."),
            ("11", "registros_insertados", "INTEGER", "No", "No", "No", "Contador total de registros efectivamente insertados en la BD externa."),
            ("12", "created_at", "DATETIME", "No", "No", "No", "Fecha y hora en que se creó este perfil de conexión."),
        ]
    },
    {
        "nombre": "comentarios",
        "descripcion": "Tabla de retroalimentación donde los usuarios pueden dejar reseñas y calificaciones sobre el sistema.",
        "objetivo": "Recolectar el feedback de los usuarios para medir la satisfacción y orientar mejoras futuras en la plataforma.",
        "relaciones": "Referencia a la tabla usuarios mediante usuario_id con eliminación en cascada.",
        "campos": [
            ("1", "id", "INTEGER", "No", "Sí", "No", "Identificador único autoincremental del comentario."),
            ("2", "usuario_id", "INTEGER", "No", "No", "Sí → usuarios.id", "ID del usuario que publicó el comentario."),
            ("3", "contenido", "VARCHAR(500)", "No", "No", "No", "Texto del comentario o reseña del usuario (máx. 500 caracteres)."),
            ("4", "calificacion", "INTEGER", "Sí", "No", "No", "Puntuación del sistema en escala del 1 al 5 (opcional)."),
            ("5", "created_at", "DATETIME", "No", "No", "No", "Fecha y hora en que se publicó el comentario."),
            ("6", "updated_at", "DATETIME", "No", "No", "No", "Fecha y hora de la última edición del comentario."),
        ]
    },
]

DDL_SCRIPT = """\
-- ============================================================
-- DICCIONARIO DE DATOS - SISTEMA GENERADOR DE DATOS SINTÉTICOS
-- Motor: SQLite / SQLAlchemy ORM
-- ============================================================

-- TABLA: usuarios
CREATE TABLE usuarios (
    id            INTEGER PRIMARY KEY AUTOINCREMENT,
    nombre        VARCHAR(100) NOT NULL,
    apellido      VARCHAR(100) NOT NULL,
    email         VARCHAR(255) NOT NULL UNIQUE,
    password_hash VARCHAR(255),
    rol           VARCHAR(20)  NOT NULL DEFAULT 'usuario',
    activo        BOOLEAN      NOT NULL DEFAULT 1,
    avatar_url    VARCHAR(500),
    created_at    DATETIME     DEFAULT (CURRENT_TIMESTAMP),
    updated_at    DATETIME     DEFAULT (CURRENT_TIMESTAMP),
    ultimo_acceso DATETIME,
    ultima_ip     VARCHAR(45)
);

-- TABLA: sesiones
CREATE TABLE sesiones (
    id           INTEGER PRIMARY KEY AUTOINCREMENT,
    usuario_id   INTEGER  NOT NULL REFERENCES usuarios(id) ON DELETE CASCADE,
    token_jwt    TEXT     NOT NULL,
    ip_address   VARCHAR(45),
    metodo_login VARCHAR(20) NOT NULL DEFAULT 'email',
    activa       BOOLEAN  NOT NULL DEFAULT 1,
    created_at   DATETIME DEFAULT (CURRENT_TIMESTAMP),
    expires_at   DATETIME NOT NULL
);

-- TABLA: logs
CREATE TABLE logs (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    usuario_id  INTEGER REFERENCES usuarios(id) ON DELETE SET NULL,
    accion      VARCHAR(100) NOT NULL,
    detalle     TEXT,
    ip_address  VARCHAR(45),
    nivel       VARCHAR(20) NOT NULL DEFAULT 'INFO',
    created_at  DATETIME DEFAULT (CURRENT_TIMESTAMP)
);

-- TABLA: conexiones
CREATE TABLE conexiones (
    id                   INTEGER PRIMARY KEY AUTOINCREMENT,
    usuario_id           INTEGER NOT NULL REFERENCES usuarios(id) ON DELETE CASCADE,
    nombre_alias         VARCHAR(100),
    motor_bd             VARCHAR(50) NOT NULL,
    host                 VARCHAR(255) NOT NULL,
    puerto               INTEGER NOT NULL,
    nombre_bd            VARCHAR(255) NOT NULL,
    usuario_db           VARCHAR(255),
    password_db          TEXT,
    registros_generados  INTEGER NOT NULL DEFAULT 0,
    registros_insertados INTEGER NOT NULL DEFAULT 0,
    created_at           DATETIME DEFAULT (CURRENT_TIMESTAMP)
);

-- TABLA: comentarios
CREATE TABLE comentarios (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    usuario_id  INTEGER NOT NULL REFERENCES usuarios(id) ON DELETE CASCADE,
    contenido   VARCHAR(500) NOT NULL,
    calificacion INTEGER,
    created_at  DATETIME DEFAULT (CURRENT_TIMESTAMP),
    updated_at  DATETIME DEFAULT (CURRENT_TIMESTAMP)
);
"""

DML_SCRIPT = """\
-- ============================================================
-- DML - INSERCIÓN DE DATOS DE EJEMPLO
-- ============================================================

-- Inserción en usuarios
INSERT INTO usuarios (nombre, apellido, email, password_hash, rol, activo)
VALUES
    ('Mariela', 'Ramos Loza', 'mariela.ramos@upt.edu.pe', 'hashed_pw_1', 'superadmin', 1),
    ('Vincenzo', 'Llanos Niño', 'vincenzo.llanos@upt.edu.pe', 'hashed_pw_2', 'usuario', 1);

-- Inserción en sesiones
INSERT INTO sesiones (usuario_id, token_jwt, ip_address, metodo_login, activa, expires_at)
VALUES
    (1, 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...', '192.168.1.10', 'email', 1, '2026-07-10 00:00:00'),
    (2, 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...', '192.168.1.11', 'google', 1, '2026-07-10 00:00:00');

-- Inserción en logs
INSERT INTO logs (usuario_id, accion, detalle, nivel)
VALUES
    (1, 'LOGIN', 'Inicio de sesión exitoso', 'INFO'),
    (2, 'GENERATE_DATA', 'Generación de 100 registros para tabla clientes', 'INFO'),
    (1, 'INJECT_DATA', 'Inyección en BD externa: mydb@localhost:3306', 'INFO');

-- Inserción en conexiones
INSERT INTO conexiones (usuario_id, nombre_alias, motor_bd, host, puerto, nombre_bd, usuario_db, password_db, registros_generados, registros_insertados)
VALUES
    (1, 'Mi MySQL Local', 'mysql', 'localhost', 3306, 'ventas_db', 'root', 'Z0Fn...cifrado...', 200, 200),
    (2, 'Postgres UPT', 'postgresql', 'db.upt.edu.pe', 5432, 'proyecto_db', 'vincenzo', 'Z0Fn...cifrado...', 150, 150);

-- Inserción en comentarios
INSERT INTO comentarios (usuario_id, contenido, calificacion)
VALUES
    (1, 'Excelente herramienta, me ahorró horas de trabajo manual generando datos de prueba.', 5),
    (2, 'Muy útil para el proyecto de BD. La integración con IA es increíble.', 5);
"""

# ──────────────────────────────────────────────────────────
# HELPERS DE FORMATO
# ──────────────────────────────────────────────────────────
AZUL_HEADER  = RGBColor(0x1F, 0x49, 0x7D)   # Azul oscuro UPT
AZUL_SUB     = RGBColor(0x27, 0x6D, 0xC3)   # Azul medio
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_CLARO   = RGBColor(0xE2, 0xEF, 0xF1)   # Para filas alternas

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def bold_run(para, text, size=11, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def normal_run(para, text, size=10):
    run = para.add_run(text)
    run.font.size = Pt(size)
    return run

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = AZUL_HEADER
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = AZUL_SUB
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = AZUL_SUB
    return p

def add_info_table(doc, tabla_info):
    """Tabla de metadata (nombre, descripcion, objetivo, relaciones)"""
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    labels = ["Nombre de la Tabla:", "Descripción de la Tabla:", "Objetivo:", "Relaciones con otras tablas:"]
    values = [tabla_info["nombre"], tabla_info["descripcion"], tabla_info["objetivo"], tabla_info["relaciones"]]
    for i, (label, value) in enumerate(zip(labels, values)):
        lc = t.rows[i].cells[0]
        vc = t.rows[i].cells[1]
        set_cell_bg(lc, "1F497D")
        lp = lc.paragraphs[0]
        r = lp.add_run(label)
        r.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(10)
        vp = vc.paragraphs[0]
        rv = vp.add_run(value)
        rv.font.size = Pt(10)
        if i == 0:
            rv.bold = True
    t.columns[0].width = Cm(5)
    t.columns[1].width = Cm(11)
    doc.add_paragraph()

def add_fields_table(doc, campos):
    """Tabla de campos de la tabla"""
    headers = ["Nro.", "Nombre del Campo", "Tipo de dato", "¿Nulos?", "PK", "FK", "Descripción del campo"]
    widths =  [Cm(1),  Cm(3.5),            Cm(2.5),        Cm(1.5),   Cm(1), Cm(3), Cm(4.5)]

    t = doc.add_table(rows=1 + len(campos), cols=len(headers))
    t.style = 'Table Grid'

    # Encabezados
    hrow = t.rows[0]
    for j, (h, w) in enumerate(zip(headers, widths)):
        c = hrow.cells[j]
        set_cell_bg(c, "276DC3")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = BLANCO
        r.font.size = Pt(9)

    # Filas de datos
    for i, campo in enumerate(campos):
        row = t.rows[i + 1]
        if i % 2 == 0:
            fill = "FFFFFF"
        else:
            fill = "E2EFF1"
        for j, val in enumerate(campo):
            c = row.cells[j]
            set_cell_bg(c, fill)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 6 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.size = Pt(9)

    doc.add_paragraph()

def add_code_block(doc, code_text):
    """Añade un bloque de código con fondo gris"""
    for line in code_text.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(8.5)
        # Fondo del párrafo
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
    doc.add_paragraph()

# ──────────────────────────────────────────────────────────
# CONSTRUCCIÓN DEL DOCUMENTO
# ──────────────────────────────────────────────────────────
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── PORTADA ──────────────────────────────────────────────
for text, bold, size, align in [
    ("UNIVERSIDAD PRIVADA DE TACNA", True,  14, WD_ALIGN_PARAGRAPH.CENTER),
    ("FACULTAD DE INGENIERÍA",        True,  13, WD_ALIGN_PARAGRAPH.CENTER),
    ("Escuela Profesional de Ingeniería de Sistemas", True, 12, WD_ALIGN_PARAGRAPH.CENTER),
    ("",                              False, 10, WD_ALIGN_PARAGRAPH.CENTER),
    ("DICCIONARIO DE DATOS",          True,  16, WD_ALIGN_PARAGRAPH.CENTER),
    ("",                              False, 10, WD_ALIGN_PARAGRAPH.CENTER),
    (f"PROYECTO: {PROYECTO}",         True,  13, WD_ALIGN_PARAGRAPH.CENTER),
    ("",                              False, 10, WD_ALIGN_PARAGRAPH.CENTER),
    (f"Curso: {CURSO}",               False, 11, WD_ALIGN_PARAGRAPH.CENTER),
    (f"Docente: {DOCENTE}",           False, 11, WD_ALIGN_PARAGRAPH.CENTER),
]:
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if text.startswith("DICCIONARIO"):
        r.font.color.rgb = AZUL_HEADER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("\n\nIntegrantes:").bold = True
for integrante in INTEGRANTES:
    pi = doc.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.add_run(integrante).font.size = Pt(11)

for text in ["", "Tacna – Perú", "2026"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(11)

doc.add_page_break()

# ── ÍNDICE GENERAL ───────────────────────────────────────
heading(doc, "ÍNDICE GENERAL", 1)
indice_items = [
    "1. Descripción de la Base de Datos",
    "2. Tablas del Sistema",
    f"   2.1. Tabla usuarios",
    f"   2.2. Tabla sesiones",
    f"   2.3. Tabla logs",
    f"   2.4. Tabla conexiones",
    f"   2.5. Tabla comentarios",
    "3. Lenguaje de Definición de Datos (DDL)",
    "4. Lenguaje de Manipulación de Datos (DML)",
]
for item in indice_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.3 if item.startswith("   ") else 0)
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# ── CUERPO ───────────────────────────────────────────────
heading(doc, "DICCIONARIO DE DATOS", 1)

heading(doc, "1. Descripción de la Base de Datos", 2)
p = doc.add_paragraph()
p.add_run(
    "El sistema utiliza una base de datos SQLite (archivo: "
).font.size = Pt(10)
r = p.add_run("datagenerator_db.db")
r.bold = True
r.font.size = Pt(10)
r.font.name = "Courier New"
p.add_run(
    ") gestionada mediante el ORM SQLAlchemy. La arquitectura de datos está organizada en 5 tablas "
    "que cubren la gestión de usuarios, autenticación, auditoría, conexiones a bases de datos externas "
    "y retroalimentación de los usuarios."
).font.size = Pt(10)

doc.add_paragraph()

heading(doc, "2. Tablas del Sistema", 2)

for idx, tabla in enumerate(TABLAS, 1):
    heading(doc, f"2.{idx}. Tabla {tabla['nombre']}", 3)
    add_info_table(doc, tabla)
    p = doc.add_paragraph()
    p.add_run("Descripción de los campos:").bold = True
    p.runs[0].font.size = Pt(10)
    add_fields_table(doc, tabla["campos"])

doc.add_page_break()

heading(doc, "3. Lenguaje de Definición de Datos (DDL)", 2)
p = doc.add_paragraph()
p.add_run(
    "A continuación se presenta el script de creación de las tablas del modelo, "
    "compatible con SQLite:"
).font.size = Pt(10)
add_code_block(doc, DDL_SCRIPT)

doc.add_page_break()

heading(doc, "4. Lenguaje de Manipulación de Datos (DML)", 2)
heading(doc, "4.1. Inserción de datos de ejemplo", 3)
p = doc.add_paragraph()
p.add_run(
    "Se muestran sentencias de inserción de ejemplo para poblar las tablas del sistema "
    "con datos de prueba representativos:"
).font.size = Pt(10)
add_code_block(doc, DML_SCRIPT)

# ── GUARDAR ──────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print("Documento guardado en: " + OUTPUT_PATH)
